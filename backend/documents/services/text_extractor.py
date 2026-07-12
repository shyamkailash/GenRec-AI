import json
from pathlib import Path

import pytesseract
from docx import Document
from PIL import Image
from pypdf import PdfReader


SUPPORTED_EXTENSIONS = {
    ".pdf",
    ".png",
    ".jpg",
    ".jpeg",
    ".docx",
    ".txt",
    ".py",
    ".ipynb",
}


class TextExtractionError(Exception):
    """Raised when text extraction fails."""


def extract_text(file_path: str) -> str:
    path = Path(file_path)

    if not path.exists():
        raise TextExtractionError(
            f"File does not exist: {path}"
        )

    extension = path.suffix.lower()

    if extension not in SUPPORTED_EXTENSIONS:
        raise TextExtractionError(
            f"Unsupported file type: {extension}"
        )

    try:
        if extension == ".pdf":
            return extract_pdf(path)

        if extension in {".png", ".jpg", ".jpeg"}:
            return extract_image(path)

        if extension == ".docx":
            return extract_docx(path)

        if extension in {".txt", ".py"}:
            return extract_plain_text(path)

        if extension == ".ipynb":
            return extract_notebook(path)

    except Exception as exc:
        raise TextExtractionError(
            f"Text extraction failed: {exc}"
        ) from exc

    return ""


def extract_pdf(path: Path) -> str:
    reader = PdfReader(str(path))

    pages = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):
        page_text = page.extract_text() or ""

        if page_text.strip():
            pages.append(
                f"--- Page {page_number} ---\n"
                f"{page_text.strip()}"
            )

    return "\n\n".join(pages).strip()


def extract_image(path: Path) -> str:
    with Image.open(path) as image:
        text = pytesseract.image_to_string(
            image,
            lang="eng",
        )

    return text.strip()


def extract_docx(path: Path) -> str:
    document = Document(str(path))

    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]

    table_content = []

    for table in document.tables:
        for row in table.rows:
            cells = [
                cell.text.strip()
                for cell in row.cells
            ]

            table_content.append(" | ".join(cells))

    content = paragraphs + table_content

    return "\n".join(content).strip()


def extract_plain_text(path: Path) -> str:
    return path.read_text(
        encoding="utf-8",
        errors="replace",
    ).strip()


def extract_notebook(path: Path) -> str:
    notebook_data = json.loads(
        path.read_text(
            encoding="utf-8",
            errors="replace",
        )
    )

    sections = []

    for cell in notebook_data.get("cells", []):
        cell_type = cell.get(
            "cell_type",
            "unknown",
        )

        source = "".join(
            cell.get("source", [])
        ).strip()

        if source:
            sections.append(
                f"[{cell_type.upper()} CELL]\n{source}"
            )

        for output in cell.get("outputs", []):
            output_text = extract_notebook_output(
                output
            )

            if output_text:
                sections.append(
                    f"[OUTPUT]\n{output_text}"
                )

    return "\n\n".join(sections).strip()


def extract_notebook_output(output: dict) -> str:
    text_content = output.get("text")

    if isinstance(text_content, list):
        return "".join(text_content).strip()

    if isinstance(text_content, str):
        return text_content.strip()

    data = output.get("data", {})
    plain_text = data.get("text/plain")

    if isinstance(plain_text, list):
        return "".join(plain_text).strip()

    if isinstance(plain_text, str):
        return plain_text.strip()

    error_name = output.get("ename")
    error_value = output.get("evalue")

    if error_name or error_value:
        return f"{error_name}: {error_value}"

    return ""
